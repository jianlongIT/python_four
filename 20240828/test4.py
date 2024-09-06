# -*- coding: utf-8 -*-
# Auther : jianlong
import os
from glob import glob
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

if __name__ == '__main__':
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.color.rgb = RGBColor(255, 51, 102)
    style.font.size = Pt(16)

    title = doc.add_heading('第一个标题', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.size = Pt(30)
    title.style.font.bold = True
    _title = title.add_run('\nhello')
    _title.italic = True
    _title.bold = True
    first_p = doc.add_paragraph('第一段内容')
    first_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    first_p.add_run('\n第二段内容')
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p2 = p2.add_run('第三段内容')
    _p2.add_picture('IMG_0326.jpg')

    table_title = ['name', 'age', 'gender']
    table_obj = doc.add_table(1, 3)
    cells = table_obj.rows[0].cells
    for i in range(3):
        cells[i].text = table_title[i]
    data = [
        ('jianlog', '22', 'boy'),
        ('jianlog', '23', 'boy'),
        ('jianlog', '23', 'boy'),
    ]
    for d in data:
        cells = table_obj.add_row().cells
        for i in range(3):
            cells[i].text = d[i]
    doc.save('test.docx')
