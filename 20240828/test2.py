# -*- coding: utf-8 -*-
# Auther : jianlong
from docx import Document
import os

if __name__ == '__main__':
    path = '/Users/dongzhiqiao/Downloads'
    docx_path = os.path.join(path, '食品安全承诺书.docx')
    result = Document(docx_path)
    p = result.paragraphs
    for i in p:
        print(i.text)

    for t in result.tables:
        for row in t.rows:
            for cell in row.cells:
                print(cell.text)
