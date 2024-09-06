# -*- coding: utf-8 -*-
# Auther : jianlong
import os
from glob import glob
from docx import Document


class MyDocument(object):
    def __init__(self, path):
        self.path = path
        self.dec = Document(path)
        self.par = ''
        self.tab = ''
        self.getPar()
        self.gettable()

    def getPar(self):
        all_par = ''
        result = self.dec.paragraphs
        for r in result:
            all_par += r.text + '\n'
        self.par = all_par

    def gettable(self):
        all_table = ''
        for t in self.dec.tables:
            for row in t.rows:
                row_text = ''
                for cell in row.cells:
                    row_text += cell.text + ','
                all_table += row_text + '\n'
        self.tab = all_table


def search(path, target):
    all_path = glob(os.path.join(path, '*'))
    path_result = []
    for p in all_path:
        if os.path.isfile(p):
            if p.endswith('.docx'):
                is_used = True
                doc = MyDocument(p)
                all_text = doc.par + doc.tab
                for t in target:
                    if t not in all_text:
                        is_used = False
                        break
                if is_used:
                    path_result.append(p)
    return path_result


if __name__ == '__main__':
    search_result = search(os.getcwd(), ['依法经营'])
    print(search_result)
